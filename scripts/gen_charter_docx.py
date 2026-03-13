"""Generate a branded DOCX of the CLARA-Gainsight Integration Charter."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import sys

# Moody's brand
NAVY = RGBColor(0x09, 0x11, 0x64)
BLUE = RGBColor(0x00, 0x5E, 0xFF)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x1F, 0x29, 0x37)
MID_GRAY = RGBColor(0x6B, 0x72, 0x80)
LIGHT_GRAY = "F2F4F8"
RED_BG = "FEE2E2"
RED_TEXT = RGBColor(0x99, 0x1B, 0x1B)
NAVY_HEX = "091164"
BLUE_HEX = "005EFF"


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val["val"]}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "auto")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_branded_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, NAVY_HEX)

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            run.font.color.rgb = DARK
            if ri % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)

    # Set column widths
    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)

    doc.add_paragraph()
    return table


def build_docx(output_path):
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # --- Default font ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = DARK

    # --- Heading styles ---
    for level in [1, 2, 3]:
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = NAVY
        if level == 1:
            hs.font.size = Pt(20)
            hs.font.bold = True
        elif level == 2:
            hs.font.size = Pt(14)
            hs.font.bold = True
        elif level == 3:
            hs.font.size = Pt(11)
            hs.font.bold = True

    # ============================================================
    # TITLE PAGE
    # ============================================================
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('CLARA-Gainsight\nIntegration Charter')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = NAVY

    doc.add_paragraph()

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run('Version 1.0 (Draft)  |  13 March 2026')
    run.font.size = Pt(14)
    run.font.color.rgb = BLUE

    doc.add_paragraph()
    doc.add_paragraph()

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_p.add_run('Author: Azmain Hossain')
    run.font.size = Pt(12)
    run.font.color.rgb = MID_GRAY
    author_p.add_run('\n')
    run2 = author_p.add_run('Customer Success Gen AI Programme')
    run2.font.size = Pt(11)
    run2.font.color.rgb = MID_GRAY
    author_p.add_run('\n')
    run3 = author_p.add_run("Moody's Analytics \u2014 Insurance Division")
    run3.font.size = Pt(11)
    run3.font.color.rgb = MID_GRAY

    doc.add_paragraph()
    status_p = doc.add_paragraph()
    status_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = status_p.add_run('DRAFT \u2014 For Review')
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    run.bold = True

    # Page break
    doc.add_page_break()

    # ============================================================
    # SECTION 1: PURPOSE
    # ============================================================
    doc.add_heading('1. Purpose', level=1)
    p = doc.add_paragraph()
    run = p.add_run('This charter defines the scope, responsibilities, timeline, and success criteria for integrating CLARA with Gainsight.')
    run.bold = True
    p.add_run('\n\nIt exists to ensure both teams have a shared understanding of what will be built, by whom, and by when \u2014 so that no party is surprised by requirements, resource demands, or delivery expectations during execution.')

    # ============================================================
    # SECTION 2: BACKGROUND
    # ============================================================
    doc.add_heading('2. Background', level=1)

    p = doc.add_paragraph()
    run = p.add_run('Gainsight')
    run.bold = True
    run.font.color.rgb = NAVY
    p.add_run(" is being rolled out as the system of record for Customer Success Managers (CSMs) across Moody's Analytics Insurance Division. It provides a unified view of customer health, engagement, time tracking, and BAU activities. CSMs are being onboarded with a production go-live target of 30 March 2026.")

    p = doc.add_paragraph()
    run = p.add_run('CLARA')
    run.bold = True
    run.font.color.rgb = NAVY
    p.add_run(' is an internally built web application that tracks IRP adoption and migration across the insurance portfolio. It serves cross-functional teams \u2014 CSMs, product, implementation, and advisory \u2014 who need visibility into IRP use cases, blockers, adoption milestones, and migration progress. Used in weekly Portfolio Reviews and primary reporting surface for senior leadership.')

    p = doc.add_paragraph()
    run = p.add_run('The problem: ')
    run.bold = True
    p.add_run('CSMs currently need to work in both systems. Without integration, CSMs must duplicate data entry, and cross-functional teams lose visibility into CSM-side activities.')

    p = doc.add_paragraph()
    run = p.add_run('The goal: ')
    run.bold = True
    p.add_run('Enable CSMs to work primarily in Gainsight while ensuring IRP adoption data flows seamlessly to CLARA, and that cross-functional inputs captured in CLARA are visible back in Gainsight.')

    # ============================================================
    # SECTION 3: PARTIES
    # ============================================================
    doc.add_heading('3. Parties & Stakeholders', level=1)

    doc.add_heading('CLARA Team', level=3)
    add_branded_table(doc,
        ['Name', 'Role'],
        [
            ['Azmain Hossain', 'Programme Manager & Lead Developer, CLARA'],
            ['Ben Brooks', 'Product Owner, CLARA'],
            ['Richard Dosoo', 'Programme & Operational Owner'],
            ['BenVH (Van Houten)', 'Infrastructure Engineer (AWS, CI/CD, App Factory)'],
        ],
        [5, 11]
    )

    doc.add_heading('Gainsight / Business Systems Team', level=3)
    add_branded_table(doc,
        ['Name', 'Role'],
        [
            ['Tina Palumbo', 'Business Systems, Gainsight Programme Lead'],
            ['Rajesh', 'Solution Architect, Gainsight'],
            ['Shashank', 'Technical Lead, Gainsight Application'],
            ['Nadeem', 'Project Manager, Gainsight Programme'],
        ],
        [5, 11]
    )

    # ============================================================
    # SECTION 4: GUIDING PRINCIPLES
    # ============================================================
    doc.add_heading('4. Guiding Principles', level=1)

    principles = [
        ('Gainsight is the system of record for CSM BAU activities.', 'CSMs should not be required to enter the same data in two systems.'),
        ('CLARA is the system of record for cross-functional IRP adoption tracking.', 'Product, implementation, and advisory teams work in CLARA and should not need Gainsight access.'),
        ('Integration must be bi-directional.', 'Gainsight-originated data must flow to CLARA. CLARA-originated data must flow back to Gainsight.'),
        ('Neither system replaces the other.', 'Gainsight covers the full customer lifecycle. CLARA covers IRP adoption specifically.'),
        ('Start small, prove connectivity, then expand.', 'A POC must demonstrate end-to-end data flow before any production integration is built.'),
        ('No disruption to active migrations.', 'The 2026 insurance scorecard target (30+ migrations) takes priority. Integration work must not divert resources.'),
    ]
    for i, (title, desc) in enumerate(principles):
        p = doc.add_paragraph()
        run = p.add_run(f'{i+1}. {title}')
        run.bold = True
        run.font.color.rgb = NAVY
        p.add_run(f'\n{desc}')

    # ============================================================
    # SECTION 5: SCOPE
    # ============================================================
    doc.add_heading('5. Scope', level=1)

    doc.add_heading('Phase 1 \u2014 Foundation (Accounts & Customer Updates)', level=2)
    add_branded_table(doc,
        ['Source Object', 'CLARA Target', 'Direction', 'Notes'],
        [
            ['Account (Parent + Sub)', 'Parent Accounts / Customers', 'GS \u2192 CLARA', 'Account hierarchy with parent-subsidiary relationships. 155+ active accounts.'],
            ['Customer Weekly Updates', 'Customer Updates', 'Bi-directional', 'Per-account summaries and meeting notes. CLARA auto-generates from transcribed calls. CSMs enter in Gainsight.'],
        ],
        [3.5, 3.5, 2.5, 6.5]
    )

    doc.add_heading('Phase 2 \u2014 Core IRP Data', level=2)
    add_branded_table(doc,
        ['Source Object', 'CLARA Target', 'Direction', 'Notes'],
        [
            ['Contact / User', 'Employees', 'GS \u2192 CLARA', 'Maps CSMs and key contacts to accounts.'],
            ['Case / Blocker', 'Blockers', 'Bi-directional', 'CSM blockers from Gainsight, product/impl blockers from CLARA.'],
            ['Success Criteria (CSC)', 'Success Criteria', 'Bi-directional', 'IRP use case success criteria tracked by CSMs and product teams.'],
        ],
        [3.5, 3.5, 2.5, 6.5]
    )

    doc.add_heading('Phase 3 \u2014 Extended Data', level=2)
    add_branded_table(doc,
        ['Source Object', 'CLARA Target', 'Direction', 'Notes'],
        [
            ['Product Adoption (PAT)', 'Product Adoption Tracking', 'GS \u2192 CLARA', 'Adoption milestone data for IRP products.'],
            ['Task / Activity', 'Action Plans / Items', 'Bi-directional', 'Cross-functional action items from both systems.'],
            ['Full Case Feed', 'Case Feed (new table)', 'GS \u2192 CLARA', 'Historical case context for blocker analysis.'],
        ],
        [3.5, 3.5, 2.5, 6.5]
    )

    doc.add_heading('Phase 4+ \u2014 Future Roadmap', level=2)
    p = doc.add_paragraph('The following are acknowledged as valuable but explicitly out of scope for the current charter. They require separate scoping once Phases 1\u20133 are delivered.')
    add_branded_table(doc,
        ['Capability', 'Description', 'Dependency'],
        [
            ['Microsoft 365 Context', 'Email + meeting signals per account', 'Graph API access, security review'],
            ['Usage Telemetry (MIDAS)', 'Product usage signals for adoption scoring', 'MIDAS integration, data pipeline'],
            ['NPS Verbatim & CTAs', 'Survey data for sentiment overlay', 'Qualtrics integration'],
            ['Gainsight Health Scores', 'CSM qualitative health (read-only)', 'Gainsight reporting API'],
            ['Write-back to SF/GS', 'Gated future capability', 'Governance framework, security review'],
        ],
        [4.5, 6, 5.5]
    )

    # ============================================================
    # SECTION 6: ARCHITECTURE
    # ============================================================
    doc.add_heading('6. Integration Architecture', level=1)

    doc.add_heading('Proposed Pattern', level=2)
    p = doc.add_paragraph("Based on the Gainsight team's presentation and technical discussion on 12 March 2026:")

    # Architecture diagram as a styled table
    arch_table = doc.add_table(rows=5, cols=3)
    arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: Salesforce
    c = arch_table.rows[0].cells[1]
    c.text = ''
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Salesforce (CRM)')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    set_cell_shading(c, LIGHT_GRAY)

    # Row 1: Arrow
    c = arch_table.rows[1].cells[1]
    c.text = ''
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\u2193')
    run.font.size = Pt(14)
    run.font.color.rgb = MID_GRAY

    # Row 2: Gainsight + CLARA side by side
    gs_cell = arch_table.rows[2].cells[0]
    gs_cell.text = ''
    p = gs_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Gainsight\n')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    run2 = p.add_run('CSM System of Record\nHealth | Engagement | Time Tracking\nUse Cases | Blockers | Activities')
    run2.font.size = Pt(8)
    run2.font.color.rgb = BLUE
    set_cell_shading(gs_cell, "EEF2FF")

    arrow_cell = arch_table.rows[2].cells[1]
    arrow_cell.text = ''
    p = arrow_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Batch (S3) \u2192\nAccounts, Contacts\n\n\u2190 API (Real-time) \u2192\nBlockers, Tasks, Updates')
    run.font.size = Pt(8)
    run.font.color.rgb = BLUE

    clara_cell = arch_table.rows[2].cells[2]
    clara_cell.text = ''
    p = clara_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('CLARA\n')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    run2 = p.add_run('Cross-Functional IRP Tracker\nMigration Reporting | Portfolio Reviews\nBlockers | Adoption | Action Plans')
    run2.font.size = Pt(8)
    run2.font.color.rgb = BLUE
    set_cell_shading(clara_cell, "EEF2FF")

    # Row 3: Databricks + App Factory
    db_cell = arch_table.rows[3].cells[0]
    db_cell.text = ''
    p = db_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Databricks\n')
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = NAVY
    run2 = p.add_run('Reporting & Analytics')
    run2.font.size = Pt(8)
    run2.font.color.rgb = MID_GRAY
    set_cell_shading(db_cell, LIGHT_GRAY)

    af_cell = arch_table.rows[3].cells[1]
    af_cell.text = ''
    p = af_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('App Factory\n')
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    run2 = p.add_run('Middleware / Orchestration')
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    set_cell_shading(af_cell, "FEF3C7")

    doc.add_paragraph()

    doc.add_heading('Integration Methods', level=2)
    add_branded_table(doc,
        ['Data Type', 'Method', 'Frequency', 'Rationale'],
        [
            ['Account (Parent + Sub)', 'Batch (S3)', 'Daily', 'Low change frequency. Proven and low-risk.'],
            ['Customer Updates', 'API (real-time)', 'On create/update', 'Auto-generated summaries. Timeliness matters.'],
            ['Contact / User', 'Batch (S3)', 'Daily', 'Low change frequency.'],
            ['Blocker', 'API (real-time)', 'On create/update', 'Immediacy needed. Delay causes duplication.'],
            ['Success Criteria', 'API (real-time)', 'On create/update', 'Active during engagements.'],
            ['Product Adoption', 'Batch (S3)', 'Daily', 'Milestone data, not time-critical.'],
            ['Task / Activity', 'API (real-time)', 'On create/update', 'Cross-functional coordination.'],
            ['Case Feed', 'Batch (S3)', 'Daily', 'Historical context.'],
        ],
        [3.5, 3, 3, 6.5]
    )

    doc.add_heading('Authentication & Connectivity', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Gainsight team will provide:')
    run.bold = True
    for item in [
        'API documentation (endpoints, schemas, rate limits, error handling)',
        'Authentication method and credentials (API key, OAuth, SSO \u2014 TBC)',
        'S3 bucket configuration for batch data exchange',
        'Sandbox / test environment access for POC',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('CLARA team will provide:')
    run.bold = True
    for item in [
        'CLARA API specification (REST endpoints for all in-scope objects)',
        "Authentication details for CLARA's API",
        'Data mapping document (Gainsight field \u2192 CLARA field + transformations)',
        'Test environment access',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('Middleware \u2014 App Factory')
    run.bold = True
    doc.add_paragraph("Integration orchestration managed through App Factory (CLARA's AWS infrastructure). Handles API orchestration, retry logic, data transformation, error logging, and rate limiting. Vendor-agnostic \u2014 can be adapted if underlying systems change.")

    # ============================================================
    # SECTION 7: RESPONSIBILITIES
    # ============================================================
    doc.add_heading('7. Responsibilities', level=1)

    doc.add_heading('Gainsight / Business Systems Team Owns:', level=3)
    for r in ['API documentation and access provisioning', 'S3 bucket setup and configuration for batch exports',
              'Sandbox environment for POC', 'Data export scheduling and reliability (batch jobs)',
              'Webhook / event configuration for real-time sync', 'Data model documentation',
              'CSM change management and training', 'Gainsight-side testing and validation', 'Project management coordination (Nadeem)']:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_heading('CLARA Team Owns:', level=3)
    for r in ['API endpoints for receiving and sending data', 'Data transformation and mapping logic',
              'Storage, indexing, and display of Gainsight-sourced data', 'Integration orchestration via App Factory',
              'CLARA-side testing and validation', 'Reporting and dashboards', 'Technical architecture and infrastructure']:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_heading('Joint Responsibilities:', level=3)
    for r in ['Data mapping agreement', 'Integration testing (end-to-end)',
              'Error handling and escalation procedures', 'POC evaluation and go/no-go decision',
              'Ongoing monitoring and incident response']:
        doc.add_paragraph(r, style='List Bullet')

    # ============================================================
    # SECTION 8: SUCCESS CRITERIA
    # ============================================================
    doc.add_heading('8. Success Criteria', level=1)

    doc.add_heading('POC Success Criteria', level=2)
    p = doc.add_paragraph('The POC is considered successful when ')
    run = p.add_run('all')
    run.bold = True
    p.add_run(' of the following are demonstrated:')

    poc = [
        ('Connectivity proven:', 'CLARA can authenticate with and read from the Gainsight API (or consume S3 exports).'),
        ('Account sync working:', 'At least 10 parent accounts with subsidiaries synced with correct hierarchy.'),
        ('Blocker round-trip:', 'A blocker created in either system appears in the other within 5 minutes.'),
        ('Data integrity:', 'Synced records match \u2014 no data loss, no field truncation, no encoding issues.'),
        ('No regression:', "CLARA's existing functionality unaffected by integration."),
        ('Performance:', 'API responses under 2 seconds. Batch jobs within agreed windows.'),
    ]
    for i, (title, desc) in enumerate(poc):
        p = doc.add_paragraph()
        run = p.add_run(f'{i+1}. {title} ')
        run.bold = True
        run.font.color.rgb = BLUE
        p.add_run(desc)

    doc.add_heading('Production Success Criteria', level=2)
    prod = [
        'All Phase 1\u20133 data objects synchronised per schedule.',
        'CSMs can enter IRP data in Gainsight and see it in CLARA without manual intervention.',
        'Cross-functional teams can enter data in CLARA and have it visible in Gainsight.',
        'Zero double-entry for objects covered by integration.',
        'CLARA reporting accurately reflects data from both sources.',
        'Integration uptime of 99.5% during business hours (08:00\u201318:00 GMT).',
    ]
    for i, item in enumerate(prod):
        p = doc.add_paragraph()
        run = p.add_run(f'{i+1}. ')
        run.bold = True
        run.font.color.rgb = BLUE
        p.add_run(item)

    # ============================================================
    # SECTION 9: TIMELINE
    # ============================================================
    doc.add_heading('9. Timeline', level=1)

    doc.add_heading('Constraints', level=2)
    constraints = [
        ('Gainsight onboarding deadline:', ' 30 March 2026 (RMS, Cape, Predicate teams). No capacity before this.'),
        ('Post-launch stabilisation:', ' ~4 weeks after go-live (April 2026) for Gainsight to stabilise.'),
        ('CLARA graduate onboarding:', ' Two graduates arriving 7 April 2026. Available from May.'),
        ('CLARA team priority:', ' 2026 scorecard target (30+ migrations) is primary. Integration is secondary.'),
    ]
    for bold, rest in constraints:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(bold)
        run.bold = True
        p.add_run(rest)

    doc.add_heading('Proposed Milestones', level=2)
    add_branded_table(doc,
        ['Phase', 'Milestone', 'Target', 'Owner'],
        [
            ['0', 'Charter signed off by both teams', 'TBD', 'Azmain / Tina'],
            ['0', 'GS API documentation shared', 'TBD', 'Rajesh / Shashank'],
            ['0', 'CLARA API spec shared', 'TBD', 'Azmain / BenVH'],
            ['0', 'Data mapping agreed', 'TBD', 'Joint'],
            ['1', 'POC environment setup', 'TBD', 'Joint'],
            ['1', 'POC: Account sync + updates', 'TBD', 'Joint'],
            ['1', 'POC: Real-time blocker sync', 'TBD', 'Joint'],
            ['1', 'POC go/no-go decision', 'TBD', 'Joint + Governance'],
            ['2', 'Phase 2 production build', 'TBD', 'Joint'],
            ['3', 'Phase 3 production build', 'TBD', 'Joint'],
            ['4+', 'Phase 4+ scoping begins', 'TBD', 'Joint'],
        ],
        [1.5, 7, 2, 5.5]
    )

    doc.add_heading('Governance Checkpoints', level=2)
    gates = [
        ('Charter Approval \u2014', ' Both teams and governance leads sign off.'),
        ('POC Go/No-Go \u2014', ' Based on POC success criteria. All six must be met.'),
        ('Phase 2 Go/No-Go \u2014', ' Based on POC results and resource availability.'),
        ('Phase 3 Go/No-Go \u2014', ' Based on Phase 2 stability and remaining scope.'),
    ]
    for i, (bold, rest) in enumerate(gates):
        p = doc.add_paragraph()
        run = p.add_run(f'Gate {i+1}: {bold}')
        run.bold = True
        run.font.color.rgb = NAVY
        p.add_run(rest)

    # ============================================================
    # SECTION 10: RISKS
    # ============================================================
    doc.add_heading('10. Risks & Mitigations', level=1)
    add_branded_table(doc,
        ['#', 'Severity', 'Risk', 'Mitigation'],
        [
            ['1', 'HIGH', 'GS API limitations prevent real-time sync', 'POC tests real-time capability. Fallback: near-real-time batch (15 min) via S3.'],
            ['2', 'HIGH', 'Integration diverts CLARA team from migrations', 'Explicitly secondary to scorecard. Governed fortnightly. Graduates from May.'],
            ['3', 'MEDIUM', 'Data model mismatch', 'Mapping agreed before build. Transformation centralised in App Factory.'],
            ['4', 'MEDIUM', 'GS team capacity constrained', 'Timeline accounts for March-April stabilisation. No POC before May.'],
            ['5', 'MEDIUM', 'Bi-directional sync creates conflicts', 'Clear ownership rules. Source system wins. Full audit trail.'],
            ['6', 'MEDIUM', 'Scope creep into Phase 4+', 'Phase 4+ deferred. Additions require governance approval + charter amendment.'],
            ['7', 'MEDIUM', 'Low CSM adoption of GS IRP workflows', 'GS team owns change management. CLARA provides cross-functional docs.'],
            ['8', 'LOW', 'Security review delays', 'Raise requests early (Phase 0). Both teams engage security in parallel.'],
        ],
        [1, 2, 5, 8]
    )

    # ============================================================
    # SECTION 11: CHANGE CONTROL
    # ============================================================
    doc.add_heading('11. Change Control', level=1)
    doc.add_paragraph('Any changes to scope, timeline, or responsibilities must:')
    steps = [
        'Be raised in writing (email or shared document comment)',
        'Be reviewed at the next fortnightly governance session',
        'Be agreed by both CLARA lead (Azmain) and Gainsight lead (Tina)',
        'Be approved by governance (Natalia Orzechowska)',
        'Be documented as a charter amendment with version and date',
    ]
    for i, st in enumerate(steps):
        p = doc.add_paragraph()
        run = p.add_run(f'{i+1}. ')
        run.bold = True
        run.font.color.rgb = BLUE
        p.add_run(st)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('No work outside the defined scope will be undertaken without an approved charter amendment.')
    run.bold = True
    run.font.color.rgb = RED_TEXT

    # ============================================================
    # SECTION 12: ASSUMPTIONS
    # ============================================================
    doc.add_heading('12. Assumptions', level=1)
    assumptions = [
        'Gainsight will have stable API and S3 export capabilities available for the POC.',
        'The Gainsight team will provide dedicated technical resource (Rajesh/Shashank).',
        "CLARA's AWS infrastructure (App Factory) can support integration middleware.",
        'Both teams will have sandbox/test environments mirroring production.',
        'CSMs will have been trained on IRP data entry in Gainsight before go-live.',
        'Salesforce data continues to flow into Gainsight and does not need independent sourcing by CLARA.',
    ]
    for i, a in enumerate(assumptions):
        p = doc.add_paragraph()
        run = p.add_run(f'{i+1}. ')
        run.bold = True
        run.font.color.rgb = BLUE
        p.add_run(a)

    # ============================================================
    # VERSION HISTORY
    # ============================================================
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('This charter is a living document.')
    run.italic = True
    run.font.color.rgb = MID_GRAY

    add_branded_table(doc,
        ['Version', 'Date', 'Author', 'Changes'],
        [['1.0', '13 Mar 2026', 'Azmain Hossain', 'Initial draft']],
        [2, 3, 4, 7]
    )

    # ============================================================
    # SAVE
    # ============================================================
    doc.save(output_path)
    print(f"DOCX created: {output_path}")


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else "generated content/CLARA_Gainsight_Integration_Charter_v1.docx"
    build_docx(out)
